from fastapi import UploadFile, HTTPException
from PyPDF2 import PdfReader
from docx import Document
import logging
from typing import Optional
import io

class TextExtractionService:
    def __init__(self):
        # Setup logging
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)
        
        # Supported file types
        self.supported_extensions = {'.pdf', '.docx', '.txt'}
        
        # Maximum file size (10MB)
        self.max_file_size = 10 * 1024 * 1024

    def extract_text(self, file: UploadFile) -> str:
        """
        Extract text from uploaded file with comprehensive error handling
        
        Args:
            file: FastAPI UploadFile object
            
        Returns:
            Extracted text content
            
        Raises:
            HTTPException: For various error conditions
        """
        if not file.filename:
            raise HTTPException(status_code=400, detail="No filename provided")
        
        filename = file.filename.lower()
        
        # Check file extension
        file_ext = self._get_file_extension(filename)
        if file_ext not in self.supported_extensions:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type. Supported types: {', '.join(self.supported_extensions)}"
            )
        
        # Check file size
        self._validate_file_size(file)
        
        try:
            # Extract text based on file type
            if file_ext == '.pdf':
                return self._extract_pdf(file)
            elif file_ext == '.docx':
                return self._extract_docx(file)
            elif file_ext == '.txt':
                return self._extract_txt(file)
            else:
                raise HTTPException(status_code=400, detail="Unsupported file type")
                
        except HTTPException:
            # Re-raise HTTP exceptions
            raise
        except Exception as e:
            self.logger.error(f"Unexpected error extracting text from {filename}: {e}")
            raise HTTPException(status_code=500, detail=f"Unexpected error processing file: {str(e)}")

    def _get_file_extension(self, filename: str) -> str:
        """Get file extension from filename"""
        if '.' not in filename:
            return ''
        return '.' + filename.split('.')[-1].lower()

    def _validate_file_size(self, file: UploadFile):
        """Validate file size"""
        # Reset file pointer to beginning
        file.file.seek(0, 2)  # Seek to end
        file_size = file.file.tell()
        file.file.seek(0)  # Reset to beginning
        
        if file_size > self.max_file_size:
            raise HTTPException(
                status_code=413, 
                detail=f"File too large. Maximum size: {self.max_file_size // (1024*1024)}MB"
            )
        
        if file_size == 0:
            raise HTTPException(status_code=400, detail="Empty file provided")

    def _extract_pdf(self, file: UploadFile) -> str:
        """Extract text from PDF file with enhanced error handling"""
        try:
            # Reset file pointer
            file.file.seek(0)
            
            # Read PDF
            reader = PdfReader(file.file)
            
            if len(reader.pages) == 0:
                raise HTTPException(status_code=400, detail="PDF file contains no pages")
            
            text_parts = []
            pages_processed = 0
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        # Clean and add page text
                        cleaned_text = self._clean_extracted_text(page_text)
                        if cleaned_text.strip():
                            text_parts.append(f"[Page {page_num + 1}]\n{cleaned_text}")
                            pages_processed += 1
                except Exception as e:
                    self.logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            if pages_processed == 0:
                raise HTTPException(status_code=400, detail="No readable text found in PDF")
            
            extracted_text = "\n\n".join(text_parts)
            self.logger.info(f"Successfully extracted text from {pages_processed} pages of PDF")
            
            return extracted_text
            
        except HTTPException:
            raise
        except Exception as e:
            self.logger.error(f"Error processing PDF: {e}")
            raise HTTPException(status_code=500, detail=f"Error processing PDF: {str(e)}")

    def _extract_docx(self, file: UploadFile) -> str:
        """Extract text from DOCX file with enhanced error handling"""
        try:
            # Reset file pointer
            file.file.seek(0)
            
            # Read DOCX
            doc = Document(file.file)
            
            text_parts = []
            paragraphs_processed = 0
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    cleaned_text = self._clean_extracted_text(para.text)
                    if cleaned_text.strip():
                        text_parts.append(cleaned_text)
                        paragraphs_processed += 1
            
            # Extract tables if any
            tables_processed = 0
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text.strip():
                    text_parts.append(f"[Table {tables_processed + 1}]\n{table_text}")
                    tables_processed += 1
            
            if paragraphs_processed == 0 and tables_processed == 0:
                raise HTTPException(status_code=400, detail="No readable text found in DOCX file")
            
            extracted_text = "\n\n".join(text_parts)
            self.logger.info(f"Successfully extracted text from DOCX: {paragraphs_processed} paragraphs, {tables_processed} tables")
            
            return extracted_text
            
        except HTTPException:
            raise
        except Exception as e:
            self.logger.error(f"Error processing DOCX: {e}")
            raise HTTPException(status_code=500, detail=f"Error processing DOCX: {str(e)}")

    def _extract_txt(self, file: UploadFile) -> str:
        """Extract text from TXT file with encoding detection"""
        try:
            # Reset file pointer
            file.file.seek(0)
            
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    file.file.seek(0)
                    content = file.file.read()
                    if isinstance(content, bytes):
                        text = content.decode(encoding)
                    else:
                        text = content
                    
                    cleaned_text = self._clean_extracted_text(text)
                    if cleaned_text.strip():
                        self.logger.info(f"Successfully extracted text from TXT file using {encoding} encoding")
                        return cleaned_text
                    else:
                        raise HTTPException(status_code=400, detail="Text file appears to be empty")
                        
                except UnicodeDecodeError:
                    continue
            
            raise HTTPException(status_code=400, detail="Could not decode text file with any supported encoding")
            
        except HTTPException:
            raise
        except Exception as e:
            self.logger.error(f"Error processing TXT file: {e}")
            raise HTTPException(status_code=500, detail=f"Error processing TXT file: {str(e)}")

    def _extract_table_text(self, table) -> str:
        """Extract text from a DOCX table"""
        table_text = []
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                table_text.append(" | ".join(row_text))
        return "\n".join(table_text)

    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = " ".join(text.split())
        
        # Remove or replace problematic characters
        text = text.replace('\x00', '')  # Remove null bytes
        text = text.replace('\r\n', '\n')  # Normalize line endings
        text = text.replace('\r', '\n')
        
        # Remove excessive newlines
        while '\n\n\n' in text:
            text = text.replace('\n\n\n', '\n\n')
        
        return text.strip()

    def get_supported_types(self) -> list:
        """Get list of supported file types"""
        return list(self.supported_extensions)

    def validate_file_type(self, filename: str) -> bool:
        """Check if file type is supported"""
        file_ext = self._get_file_extension(filename)
        return file_ext in self.supported_extensions