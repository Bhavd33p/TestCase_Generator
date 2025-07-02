from typing import List
import logging
from fastapi import UploadFile, HTTPException
from langchain_core.documents import Document
from PyPDF2 import PdfReader
import docx
import io

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

async def process_uploaded_file(file: UploadFile, file_type: str) -> List[Document]:
    """
    Process an uploaded file and extract text.
    
    Args:
        file (UploadFile): The uploaded file
        file_type (str): Type of file (pdf, word, etc.)
        
    Returns:
        List[Document]: List of documents with extracted text
    """
    try:
        # Read file content
        content = await file.read()
        
        # Extract text based on file type
        if file_type.lower() == 'pdf':
            text = extract_text_from_pdf(content)
        elif file_type.lower() in ['docx', 'word']:
            text = extract_text_from_docx(content)
        elif file_type.lower() == 'txt':
            text = extract_text_from_txt(content)
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {file_type}. Supported types are: pdf, docx, txt"
            )
        
        # Create document with metadata
        document = Document(
            page_content=text,
            metadata={
                "filename": file.filename,
                "file_type": file_type,
                "content_type": file.content_type
            }
        )
        
        return [document]
        
    except Exception as e:
        logger.error(f"Error processing file: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to process file: {str(e)}"
        )

def extract_text_from_pdf(content: bytes) -> str:
    """
    Extract text from PDF content.
    
    Args:
        content (bytes): PDF file content
        
    Returns:
        str: Extracted text
    """
    try:
        pdf_file = io.BytesIO(content)
        pdf_reader = PdfReader(pdf_file)
        text = ""
        
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
            
        return text.strip()
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise Exception(f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_docx(content: bytes) -> str:
    """
    Extract text from DOCX content.
    
    Args:
        content (bytes): DOCX file content
        
    Returns:
        str: Extracted text
    """
    try:
        docx_file = io.BytesIO(content)
        doc = docx.Document(docx_file)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
            
        return text.strip()
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")

def extract_text_from_txt(content: bytes) -> str:
    """
    Extract text from TXT content.
    
    Args:
        content (bytes): TXT file content
        
    Returns:
        str: Extracted text
    """
    try:
        return content.decode('utf-8').strip()
        
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {str(e)}")
        raise Exception(f"Failed to extract text from TXT: {str(e)}") 