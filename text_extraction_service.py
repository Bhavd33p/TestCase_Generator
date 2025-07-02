import logging
import os
import tempfile
from typing import Dict, Any, List, Optional, Union
from pathlib import Path
import mimetypes
import chardet

# Document processing imports
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class TextExtractionService:
    """
    Comprehensive text extraction service supporting multiple file formats
    with robust error handling and validation.
    """
    
    def __init__(self, max_file_size: int = 50 * 1024 * 1024):  # 50MB default
        """
        Initialize the text extraction service.
        
        Args:
            max_file_size: Maximum file size in bytes (default: 50MB)
        """
        self.max_file_size = max_file_size
        self.supported_extensions = {'.txt', '.pdf', '.docx'}
        
        # Check available libraries
        self.pdf_available = PDF_AVAILABLE
        self.docx_available = DOCX_AVAILABLE
        
        if not self.pdf_available:
            logger.warning("PyPDF2 not available - PDF extraction disabled")
        if not self.docx_available:
            logger.warning("python-docx not available - DOCX extraction disabled")
        
        logger.info(f"TextExtractionService initialized with max file size: {max_file_size} bytes")
    
    def _validate_file(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Validate file before processing.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Validation result dictionary
        """
        file_path = Path(file_path)
        
        # Check if file exists
        if not file_path.exists():
            return {
                "valid": False,
                "error": f"File does not exist: {file_path}",
                "error_type": "file_not_found"
            }
        
        # Check if it's a file (not directory)
        if not file_path.is_file():
            return {
                "valid": False,
                "error": f"Path is not a file: {file_path}",
                "error_type": "not_a_file"
            }
        
        # Check file size
        file_size = file_path.stat().st_size
        if file_size > self.max_file_size:
            return {
                "valid": False,
                "error": f"File too large: {file_size} bytes (max: {self.max_file_size})",
                "error_type": "file_too_large"
            }
        
        if file_size == 0:
            return {
                "valid": False,
                "error": "File is empty",
                "error_type": "empty_file"
            }
        
        # Check file extension
        extension = file_path.suffix.lower()
        if extension not in self.supported_extensions:
            return {
                "valid": False,
                "error": f"Unsupported file type: {extension}. Supported: {self.supported_extensions}",
                "error_type": "unsupported_format"
            }
        
        # Check library availability for specific formats
        if extension == '.pdf' and not self.pdf_available:
            return {
                "valid": False,
                "error": "PDF processing not available (PyPDF2 not installed)",
                "error_type": "library_unavailable"
            }
        
        if extension == '.docx' and not self.docx_available:
            return {
                "valid": False,
                "error": "DOCX processing not available (python-docx not installed)",
                "error_type": "library_unavailable"
            }
        
        return {
            "valid": True,
            "file_size": file_size,
            "extension": extension,
            "mime_type": mimetypes.guess_type(str(file_path))[0]
        }
    
    def _detect_encoding(self, file_path: Path) -> str:
        """
        Detect file encoding for text files.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Detected encoding or 'utf-8' as fallback
        """
        try:
            with open(file_path, 'rb') as file:
                raw_data = file.read(10000)  # Read first 10KB for detection
                result = chardet.detect(raw_data)
                encoding = result.get('encoding', 'utf-8')
                confidence = result.get('confidence', 0)
                
                logger.debug(f"Detected encoding: {encoding} (confidence: {confidence:.2f})")
                
                # Use utf-8 if confidence is too low
                if confidence < 0.7:
                    logger.warning(f"Low confidence in encoding detection, using utf-8")
                    return 'utf-8'
                
                return encoding or 'utf-8'
                
        except Exception as e:
            logger.warning(f"Failed to detect encoding: {e}, using utf-8")
            return 'utf-8'
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw text to clean
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = " ".join(text.split())
        
        # Remove control characters except newlines and tabs
        cleaned_chars = []
        for char in text:
            if ord(char) >= 32 or char in ['\n', '\t']:
                cleaned_chars.append(char)
        
        text = "".join(cleaned_chars)
        
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove excessive newlines
        while '\n\n\n' in text:
            text = text.replace('\n\n\n', '\n\n')
        
        return text.strip()
    
    def _extract_text_file(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text from plain text files.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Extraction result
        """
        try:
            encoding = self._detect_encoding(file_path)
            
            with open(file_path, 'r', encoding=encoding, errors='replace') as file:
                content = file.read()
            
            cleaned_content = self._clean_text(content)
            
            return {
                "success": True,
                "content": cleaned_content,
                "metadata": {
                    "file_type": "text",
                    "encoding": encoding,
                    "character_count": len(cleaned_content),
                    "line_count": cleaned_content.count('\n') + 1 if cleaned_content else 0
                }
            }
            
        except UnicodeDecodeError as e:
            logger.error(f"Unicode decode error for {file_path}: {e}")
            return {
                "success": False,
                "error": f"Failed to decode text file: {e}",
                "error_type": "encoding_error"
            }
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            return {
                "success": False,
                "error": f"Text extraction failed: {e}",
                "error_type": "extraction_error"
            }
    
    def _extract_pdf_file(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text from PDF files.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extraction result
        """
        if not self.pdf_available:
            return {
                "success": False,
                "error": "PDF processing not available (PyPDF2 not installed)",
                "error_type": "library_unavailable"
            }
        
        try:
            content_parts = []
            page_count = 0
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            content_parts.append(page_text)
                        logger.debug(f"Extracted text from page {page_num + 1}")
                    except Exception as page_error:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {page_error}")
                        continue
            
            # Combine all pages
            full_content = "\n\n".join(content_parts)
            cleaned_content = self._clean_text(full_content)
            
            if not cleaned_content:
                return {
                    "success": False,
                    "error": "No text content found in PDF",
                    "error_type": "no_content"
                }
            
            return {
                "success": True,
                "content": cleaned_content,
                "metadata": {
                    "file_type": "pdf",
                    "page_count": page_count,
                    "pages_processed": len(content_parts),
                    "character_count": len(cleaned_content)
                }
            }
            
        except Exception as e:
            logger.error(f"Failed to extract PDF {file_path}: {e}")
            return {
                "success": False,
                "error": f"PDF extraction failed: {e}",
                "error_type": "extraction_error"
            }
    
    def _extract_docx_file(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text from DOCX files.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extraction result
        """
        if not self.docx_available:
            return {
                "success": False,
                "error": "DOCX processing not available (python-docx not installed)",
                "error_type": "library_unavailable"
            }
        
        try:
            doc = Document(file_path)
            content_parts = []
            
            # Extract text from paragraphs
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)
                    paragraph_count += 1
            
            # Extract text from tables
            table_count = 0
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    content_parts.append("\n".join(table_text))
                    table_count += 1
            
            # Combine all content
            full_content = "\n\n".join(content_parts)
            cleaned_content = self._clean_text(full_content)
            
            if not cleaned_content:
                return {
                    "success": False,
                    "error": "No text content found in DOCX",
                    "error_type": "no_content"
                }
            
            return {
                "success": True,
                "content": cleaned_content,
                "metadata": {
                    "file_type": "docx",
                    "paragraph_count": paragraph_count,
                    "table_count": table_count,
                    "character_count": len(cleaned_content)
                }
            }
            
        except Exception as e:
            logger.error(f"Failed to extract DOCX {file_path}: {e}")
            return {
                "success": False,
                "error": f"DOCX extraction failed: {e}",
                "error_type": "extraction_error"
            }
    
    def extract_text(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Extract text from a file.
        
        Args:
            file_path: Path to the file to extract text from
            
        Returns:
            Dictionary containing extraction results
        """
        file_path = Path(file_path)
        
        # Validate file
        validation = self._validate_file(file_path)
        if not validation["valid"]:
            logger.error(f"File validation failed: {validation['error']}")
            return {
                "success": False,
                "file_path": str(file_path),
                "error": validation["error"],
                "error_type": validation["error_type"]
            }
        
        extension = validation["extension"]
        logger.info(f"Extracting text from {extension} file: {file_path}")
        
        # Route to appropriate extraction method
        try:
            if extension == '.txt':
                result = self._extract_text_file(file_path)
            elif extension == '.pdf':
                result = self._extract_pdf_file(file_path)
            elif extension == '.docx':
                result = self._extract_docx_file(file_path)
            else:
                return {
                    "success": False,
                    "file_path": str(file_path),
                    "error": f"Unsupported file type: {extension}",
                    "error_type": "unsupported_format"
                }
            
            # Add common metadata
            if result["success"]:
                result["file_path"] = str(file_path)
                result["file_size"] = validation["file_size"]
                result["mime_type"] = validation.get("mime_type")
                
                logger.info(f"Successfully extracted {len(result['content'])} characters from {file_path}")
            else:
                result["file_path"] = str(file_path)
                logger.error(f"Failed to extract text from {file_path}: {result.get('error', 'Unknown error')}")
            
            return result
            
        except Exception as e:
            logger.error(f"Unexpected error during text extraction from {file_path}: {e}")
            return {
                "success": False,
                "file_path": str(file_path),
                "error": f"Unexpected extraction error: {e}",
                "error_type": "unexpected_error"
            }
    
    def extract_text_from_bytes(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract text from file bytes (useful for uploaded files).
        
        Args:
            file_bytes: File content as bytes
            filename: Original filename (used to determine file type)
            
        Returns:
            Dictionary containing extraction results
        """
        if not file_bytes:
            return {
                "success": False,
                "error": "Empty file bytes provided",
                "error_type": "empty_input"
            }
        
        if len(file_bytes) > self.max_file_size:
            return {
                "success": False,
                "error": f"File too large: {len(file_bytes)} bytes (max: {self.max_file_size})",
                "error_type": "file_too_large"
            }
        
        # Create temporary file
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as temp_file:
                temp_file.write(file_bytes)
                temp_file_path = temp_file.name
            
            # Extract text from temporary file
            result = self.extract_text(temp_file_path)
            
            # Add original filename to metadata
            if result["success"] and "metadata" in result:
                result["metadata"]["original_filename"] = filename
            
            return result
            
        except Exception as e:
            logger.error(f"Failed to process file bytes for {filename}: {e}")
            return {
                "success": False,
                "error": f"Failed to process file bytes: {e}",
                "error_type": "processing_error"
            }
        finally:
            # Clean up temporary file
            try:
                if 'temp_file_path' in locals():
                    os.unlink(temp_file_path)
            except Exception:
                pass
    
    def get_supported_formats(self) -> Dict[str, Any]:
        """
        Get information about supported file formats.
        
        Returns:
            Dictionary with supported formats and their status
        """
        return {
            "supported_extensions": list(self.supported_extensions),
            "format_status": {
                "txt": {"available": True, "description": "Plain text files"},
                "pdf": {
                    "available": self.pdf_available,
                    "description": "PDF documents",
                    "library": "PyPDF2"
                },
                "docx": {
                    "available": self.docx_available,
                    "description": "Microsoft Word documents",
                    "library": "python-docx"
                }
            },
            "max_file_size": self.max_file_size,
            "max_file_size_mb": self.max_file_size / (1024 * 1024)
        }
    
    def validate_file_upload(self, filename: str, file_size: int) -> Dict[str, Any]:
        """
        Validate a file upload before processing.
        
        Args:
            filename: Name of the uploaded file
            file_size: Size of the uploaded file in bytes
            
        Returns:
            Validation result
        """
        extension = Path(filename).suffix.lower()
        
        if extension not in self.supported_extensions:
            return {
                "valid": False,
                "error": f"Unsupported file type: {extension}",
                "error_type": "unsupported_format"
            }
        
        if file_size > self.max_file_size:
            return {
                "valid": False,
                "error": f"File too large: {file_size} bytes (max: {self.max_file_size})",
                "error_type": "file_too_large"
            }
        
        if file_size == 0:
            return {
                "valid": False,
                "error": "File is empty",
                "error_type": "empty_file"
            }
        
        # Check library availability
        if extension == '.pdf' and not self.pdf_available:
            return {
                "valid": False,
                "error": "PDF processing not available",
                "error_type": "library_unavailable"
            }
        
        if extension == '.docx' and not self.docx_available:
            return {
                "valid": False,
                "error": "DOCX processing not available",
                "error_type": "library_unavailable"
            }
        
        return {
            "valid": True,
            "extension": extension,
            "file_size": file_size
        }
    
    def health_check(self) -> Dict[str, Any]:
        """
        Perform health check on the text extraction service.
        
        Returns:
            Health status information
        """
        return {
            "status": "healthy",
            "supported_formats": len(self.supported_extensions),
            "libraries": {
                "PyPDF2": self.pdf_available,
                "python-docx": self.docx_available
            },
            "max_file_size_mb": self.max_file_size / (1024 * 1024)
        } 